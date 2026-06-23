#!/usr/bin/env python3
"""Generate IMP SSO implementation Word document with embedded SVG diagrams."""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from reportlab.graphics import renderPM
from svglib.svglib import svg2rlg

ROOT = Path(__file__).resolve().parent
DIAGRAMS = ROOT / "diagrams"
OUTPUT_DOCX = ROOT / "IMP_SSO_实现方案.docx"


def svg_to_png(svg_path: Path, png_path: Path, scale: float = 2.0) -> None:
    drawing = svg2rlg(str(svg_path))
    if drawing is None:
        raise RuntimeError(f"Failed to parse SVG: {svg_path}")
    renderPM.drawToFile(drawing, str(png_path), fmt="PNG", dpi=72 * scale)


def set_doc_default_font(document: Document, font_name: str = "Microsoft YaHei", size: int = 11) -> None:
    style = document.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(size)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def add_title(document: Document, text: str) -> None:
    p = document.add_heading(text, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_heading(document: Document, text: str, level: int = 1) -> None:
    document.add_heading(text, level=level)


def add_paragraph(document: Document, text: str, bold: bool = False) -> None:
    p = document.add_paragraph()
    run = p.add_run(text)
    run.bold = bold


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = header
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value


def add_diagram(document: Document, title: str, svg_name: str, caption: str) -> None:
    svg_path = DIAGRAMS / svg_name
    png_path = DIAGRAMS / svg_name.replace(".svg", ".png")
    svg_to_png(svg_path, png_path)

    add_heading(document, title, level=2)
    document.add_picture(str(png_path), width=Inches(6.5))
    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    note = document.add_paragraph()
    note.add_run(f"SVG 源文件：diagrams/{svg_name}").italic = True


def build_document() -> Document:
    document = Document()
    set_doc_default_font(document)

    add_title(document, "Athlon Agent IMP SSO 实现方案")
    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"文档版本：V2.0    创建日期：{date.today().isoformat()}\n").bold = True
    meta.add_run("适用范围：Athlon Agent（feature/imp-sso）")

    add_heading(document, "1. 概述")
    add_paragraph(
        document,
        "Athlon Agent 作为企业级 WPF 桌面应用，通过 IMP（ICBC Asia）SSO 实现统一身份认证。"
        "本实现严格遵循 IMP/SsoFilter 实际协议（ssotoken + check_ssotoken），不采用 OAuth 2.0 模型。",
    )
    add_bullets(
        document,
        [
            "默认 Sso.Enabled = false，前期不影响现有 License 流程",
            "启用后：SSO 与 License 双重门控均须通过",
            "本地会话固定 24 小时有效（SessionValidityHours），不做续期",
            "标题栏显示 IMP 返回的 ename，支持登出",
            "浏览器回调采用双端点设计：GET /sso/auth 返回认证页，POST /sso/complete 回传 token",
        ],
    )

    add_heading(document, "2. 设计原则")
    add_table(
        document,
        ["不使用", "使用（IMP 实际）"],
        [
            ["OAuth access/refresh token", "单一 ssotoken（回调 hash 参数 token）"],
            ["POST /refresh_token", "无 refresh；登录时 POST check_ssotoken 一次"],
            ["IMP timeoutRemaining 驱动本地过期", "本地 LoggedInAt + SessionValidityHours"],
            ["SsoFilter 5 分钟轮询续期", "不实现续期"],
            ["服务端直接读取 hash", "浏览器 JS 解析 hash 后 POST JSON 到 /sso/complete"],
        ],
    )

    add_heading(document, "3. 配置说明")
    add_paragraph(document, "settings.json 示例（~/.athlon-agent/config/settings.json）：")
    document.add_paragraph(
        '{\n'
        '  "Sso": {\n'
        '    "Enabled": false,\n'
        '    "ImpDomain": "imp.icbcasiauat.com",\n'
        '    "AppId": "252",\n'
        '    "Msg": "123456789",\n'
        '    "Version": "20251127",\n'
        '    "SessionValidityHours": 24,\n'
        '    "CallbackPort": 5657,\n'
        '    "CallbackPath": "/sso/auth",\n'
        '    "CompletePath": "/sso/complete"\n'
        '  }\n'
        '}',
        style="Intense Quote",
    )
    add_table(
        document,
        ["配置项", "说明"],
        [
            ["Enabled", "是否启用 IMP SSO，默认 false"],
            ["ImpDomain", "IMP 域名，UAT 默认 imp.icbcasiauat.com，生产切换为正式域名"],
            ["AppId", "IMP 子应用 ID（impappid）"],
            ["Msg", "IMP 子应用标识，用于 index.html 的 msg 参数"],
            ["Version", "check_ssotoken 接口版本号"],
            ["SessionValidityHours", "本地会话有效期（小时），默认 24"],
            ["CallbackPort", "本地回调监听端口，默认 5657"],
            ["CallbackPath", "认证页路径，默认 /sso/auth"],
            ["CompletePath", "token 回传路径，默认 /sso/complete"],
        ],
    )
    add_table(
        document,
        ["Sso.Enabled", "行为"],
        [
            ["false（默认）", "跳过 IMP；仅 License；无 SSO UI"],
            ["true", "SSO 通过后再 License；显示 ename；可登出"],
        ],
    )

    add_heading(document, "4. Token 生命周期")
    add_bullets(
        document,
        [
            "IMP 登录成功 → 浏览器 POST /sso/complete → Agent POST check_ssotoken（仅此一次）→ 存 ssotoken + ename",
            "ExpiresAt = LoggedInAt + SessionValidityHours（默认 24h）",
            "会话持久化至 ~/.athlon-agent/credentials/sso-session.secret（DPAPI 加密）",
            "24h 内再次启动：读 DPAPI 缓存，不调 IMP、不打开浏览器",
            "超过 24h：清缓存 → 重新浏览器 IMP 登录",
            "运行期间不续期、不轮询 check_ssotoken",
        ],
    )

    add_diagram(
        document,
        "4.1 启动门控流程图",
        "startup_flow.svg",
        "图 1  应用启动时 SSO / License 门控流程",
    )

    add_diagram(
        document,
        "4.2 IMP 登录时序图",
        "login_sequence.svg",
        "图 2  IMP 浏览器登录与双端点回调时序",
    )

    add_heading(document, "5. IMP 协议要点")
    add_heading(document, "5.1 端点", level=2)
    add_table(
        document,
        ["端点", "何时调用"],
        [
            ["/icbcasia/imp/index.html?toLogin=true&appId={AppId}&msg={Msg}#/login", "无有效缓存或已过期时打开浏览器"],
            ["/icbcasia/imp/check_ssotoken", "收到 /sso/complete 回调后调用一次"],
            ["GET localhost:{CallbackPort}/sso/auth", "IMP 重定向后浏览器请求认证页"],
            ["POST localhost:{CallbackPort}/sso/complete", "认证页 JS 解析 hash 后回传 JSON"],
        ],
    )

    add_heading(document, "5.2 check_ssotoken 请求", level=2)
    document.add_paragraph(
        "POST https://{ImpDomain}/icbcasia/imp/check_ssotoken?version={Version}\n"
        "Content-Type: application/x-www-form-urlencoded\n\n"
        "ssotoken={token}&impappid={AppId}&dse_sessionId={token}&requestip={localIp}",
        style="Intense Quote",
    )

    add_heading(document, "5.3 check_ssotoken 响应解析", level=2)
    add_table(
        document,
        ["条件", "状态", "处理"],
        [
            ["retcode = 1", "ReLoginRequired", "提示重新登录"],
            ["roleUserRelNum <= 0", "NoRole", "提示无角色，打开 noRoleForSubApp 页面"],
            ["返回 HTML 登录页", "LoginRequired", "需要 IMP 登录"],
            ["userid + timeoutRemaining > 0", "Valid", "提取 ename 作为 DisplayName，创建会话"],
            ["HTTP 非 200 或解析失败", "Invalid", "显示错误信息，退出应用"],
        ],
    )

    add_heading(document, "5.4 回调 URL 与双端点流程", level=2)
    add_paragraph(document, "IMP 登录成功后重定向至：")
    document.add_paragraph(
        "http://localhost:5657/sso/auth#/?appId=252&userId=000974115&locale=zh_HK"
        "&token={ssotoken}&role=User&depname=FTD",
        style="Intense Quote",
    )
    add_bullets(
        document,
        [
            "token 在 URL hash 中，HttpListener 无法直接读取",
            "GET /sso/auth → 返回 ImpSsoAuthPageHtml 认证页（含加载/成功/失败状态）",
            "浏览器 JS 解析 hash → POST /sso/complete（JSON body），连接保持挂起",
            "Agent 收到 payload 后校验 appId，调用 check_ssotoken 二次校验",
            "验票完成后 CompleteBrowserResponseAsync 返回 {ok, message}，浏览器与桌面端状态一致",
            "重复 POST 返回 409；登录流程结束后由调用方显式 StopAsync 关闭监听器",
        ],
    )

    add_heading(document, "5.5 /sso/complete 请求体", level=2)
    document.add_paragraph(
        '{\n'
        '  "appId": "252",\n'
        '  "userId": "000974115",\n'
        '  "locale": "zh_HK",\n'
        '  "token": "{ssotoken}",\n'
        '  "role": "User",\n'
        '  "depname": "FTD",\n'
        '  "channel_type": null,\n'
        '  "msg": null\n'
        '}',
        style="Intense Quote",
    )

    add_heading(document, "6. 核心组件")
    add_table(
        document,
        ["组件", "路径", "职责"],
        [
            ["SsoSettings", "Athlon.Agent.Core/Sso/", "配置项，含 Msg、Enabled 默认 false"],
            ["ImpSsoSession", "Athlon.Agent.Core/Sso/", "SsoToken、UserId、DisplayName、LoggedInAt、ExpiresAt"],
            ["ImpSsoCallbackPayload", "Athlon.Agent.Core/Sso/", "浏览器回传的 hash 参数字段"],
            ["ImpSsoSessionStore", "Infrastructure/Sso/", "DPAPI 加密持久化至 credentials/sso-session.secret"],
            ["ImpSsoAuthService", "Infrastructure/Sso/", "构建登录 URL、POST check_ssotoken、CompleteLoginAsync"],
            ["ImpSsoResponseParser", "Infrastructure/Sso/", "解析 check_ssotoken JSON 响应"],
            ["ImpSsoCallbackServer", "Infrastructure/Sso/", "5657 端口双端点回调服务"],
            ["ImpSsoAuthPageHtml", "Infrastructure/Sso/", "认证页 HTML（hash 解析 + fetch POST）"],
            ["ImpSsoStartupGate", "App/Licensing/", "启动门控编排、等待窗口、浏览器登录"],
            ["ImpSsoLoginWaitingWindow", "App/Licensing/", "登录等待提示窗口"],
            ["MainWindowViewModel", "App/ViewModels/", "标题栏 ename 显示 + 登出"],
        ],
    )

    add_heading(document, "7. 启动顺序")
    document.add_paragraph(
        "1. StartupUpdateGate.CheckBeforeStartupGates\n"
        "2. 若 Sso.Enabled → ImpSsoStartupGate.EnsureAuthenticated\n"
        "   a. 读 DPAPI 缓存，未过期则直接通过\n"
        "   b. 过期则显示 ImpSsoLoginWaitingWindow，后台线程执行浏览器登录\n"
        "   c. 登录超时 5 分钟\n"
        "3. LicenseStartupGate.EnsureLicensed\n"
        "4. DI 注册 + MainWindow",
        style="List Number",
    )

    add_heading(document, "8. UI 行为")
    add_bullets(
        document,
        [
            "Sso.Enabled=true：标题栏右侧显示 ename（IMP 返回，无 ename 时回退 userid）",
            "点击姓名 → 确认「确定要退出登录吗？」→ 清 ssotoken → 退出应用",
            "Sso.Enabled=false：不显示 SSO 相关 UI",
            "浏览器认证页：加载中 → 成功（可关闭）/ 失败（显示 IMP msg 或通用错误）",
            "登录等待窗口：提示用户已在浏览器中打开登录页，登录成功后自动关闭",
        ],
    )

    add_heading(document, "9. 安全设计")
    add_bullets(
        document,
        [
            "ssotoken 使用 Windows DPAPI 加密存储（CurrentUser），文件路径 credentials/sso-session.secret",
            "回调 token 不可信，必须 check_ssotoken 二次校验",
            "回调 appId 与配置 AppId 不一致时拒绝登录",
            "日志中应对 ssotoken 脱敏",
            "全部 IMP 通信使用 HTTPS",
            "CallbackServer 在收到回调后优雅关闭，等待进行中的请求排空（最多 5 秒）",
        ],
    )

    add_heading(document, "10. 错误处理")
    add_table(
        document,
        ["场景", "用户提示", "应用行为"],
        [
            ["登录超时（5 分钟）", "IMP 登录超时，请重试", "退出应用"],
            ["NoRole（无可用角色）", "当前登录用户无可用角色", "打开 noRoleForSubApp 页面，退出应用"],
            ["ReLoginRequired", "IMP 返回的 retmsg", "退出应用"],
            ["hash 无 token", "浏览器页显示 IMP msg 或通用错误", "等待超时后退出"],
            ["check_ssotoken 失败", "IMP SSO 登录失败", "退出应用"],
        ],
    )

    add_heading(document, "11. 验收标准")
    add_heading(document, "11.1 Sso.Enabled=false（默认）", level=2)
    add_bullets(
        document,
        [
            "启动时不打开浏览器、不监听 5657、不调 IMP",
            "仅走 License 门控",
            "标题栏无 SSO 姓名/登出入口",
        ],
    )
    add_heading(document, "11.2 Sso.Enabled=true", level=2)
    add_bullets(
        document,
        [
            "首次 IMP 登录 → GET /sso/auth → POST /sso/complete（挂起）→ check_ssotoken → 返回最终 JSON → 进入主界面",
            "浏览器认证页在 IMP 验票完成后显示成功或失败状态",
            "24h 内再次启动免浏览器、免 IMP",
            "超过 24h 重新 IMP 登录",
            "SSO + License 双门控",
            "点击姓名可登出",
            "无角色用户引导至 noRoleForSubApp 页面",
        ],
    )

    add_heading(document, "12. 附录：Debug 开关")
    add_paragraph(
        document,
        "DEBUG 构建可设置环境变量 ATHLON_SKIP_SSO=1 跳过 SSO 门控（与 ATHLON_SKIP_LICENSE 对称）。",
    )

    return document


def main() -> None:
    DIAGRAMS.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(OUTPUT_DOCX)
    print(f"Generated: {OUTPUT_DOCX}")
    print(f"SVG diagrams: {DIAGRAMS / 'startup_flow.svg'}, {DIAGRAMS / 'login_sequence.svg'}")


if __name__ == "__main__":
    main()
